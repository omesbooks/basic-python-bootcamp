#GUIWiki.py
import wikipedia
from docx import Document

def Wiki(keyword,lang='th'):
       wikipedia.set_lang(lang)
       #summary สำหรับบทความที่สรุป
       data = wikipedia.summary(keyword)
       
       #page+content บทความทั้งหน้า
       data2 = wikipedia.page(keyword)
       data2 = data2.content

       doc = Document() # สร้างไฟล์ word ใน python
       doc.add_heading(keyword,0)

       doc.add_paragraph(data2) #เอาข้อมูลที่ดึงจากวิกิมาใส่เป็น paragraph
       doc.save(keyword+'.docx')
       print('สร้างไฟล์สำเร็จ')       
           
#เปลี่ยนภาษาไทย
wikipedia.set_lang('th')
from tkinter import *
from tkinter import ttk
from tkinter import messagebox

GUI = Tk()
GUI.title('โปรแกรม wiki')
GUI.geometry('400x300')
# config
FONT1 = ('Angsana New',15)

#คำอธิบาย
L = ttk.Label(GUI, text = 'ค้นหาบทความ',font=FONT1)
L.pack()

#ช่องค้นหาข้อมูล
v_search =  StringVar() #กล่องสำหรับเก็บ keyword
E1 = ttk.Entry(GUI,textvariable=v_search,font = FONT1,width=40)
E1.pack(pady = 10)



#ปุ่มค้นหา
def Search():
       keyword = v_search.get() # .get() คือการดึงข้อมูลเข้ามา
       try:                                     #ทดลองว่าการ search ได้ผลลัพธ์หรือไม่
           language = v_radio.get()
           Wiki(keyword,language)
           messagebox.showwarning('บันทึกไฟล์สำเร็จ','ค้นหาข้อความสำเร็จ')
       except:                                  #หากรันแล้วมีปัญหาให้แสดงข้อความแจ้งเตือน
            messagebox.showwarning('keyword error','กรุณากรอกคำค้นหาใหม่')
            
            
       '''  
       print(wikipedia.search(keyword))
       result = wikipedia.summary(keyword)
       print(result)       
       ''' 
B1 = ttk.Button(GUI,text='Search',command = Search)
B1.pack(ipadx=20,ipady=10) #ipadx ขยายปุ่มในแนวแกน x


#เลือกภาษา
F1 = Frame(GUI)
F1.pack(pady=10)

v_radio = StringVar() #ตัวแปรที่เอาไว้เก็บข้อมูล

RB1 = ttk.Radiobutton(F1,text='ภาษาไทย',variable=v_radio,value='th')
RB2 = ttk.Radiobutton(F1,text='อังกฤษ',variable=v_radio,value='en')
RB3 = ttk.Radiobutton(F1,text='จีน',variable=v_radio,value='zh')
RB1.invoke() #สั่งให้ค่าเริ่มต้นเป็นภาษาไทย

RB1.grid(row=0,column=0)
RB2.grid(row=0,column=1)
RB3.grid(row=0,column=2)



GUI.mainloop()
