#test-docx.py
from docx import Document
import wikipedia as wk      #ไม่มี from 

def Wiki(keyword_1,lang='th'):
       wk.set_lang(lang)
       #summary สำหรับบทความที่สรุป
       data = wk.summary(keyword_1)
       
       #page+content บทความทั้งหน้า
       data2 = wk.page(keyword_1)
       data2 = data2.content

       doc = Document() # สร้างไฟล์ word ใน python
       doc.add_heading(keyword_1,0)

       doc.add_paragraph(data2) #เอาข้อมูลที่ดึงจากวิกิมาใส่เป็น paragraph
       doc.save(keyword_1+'.docx')
       print('สร้างไฟล์สำเร็จ')
       
Wiki('KFC')
Wiki('Texas Chicken','en')
Wiki('Mcdonald','jp')
